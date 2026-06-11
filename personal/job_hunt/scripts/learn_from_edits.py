#!/usr/bin/env python3
"""
learn_from_edits.py — Feedback loop: learn from document edits.

Scans applications/ for pairs of:
  - *_rough_draft.docx  (LLM-generated output)
  - *.docx without the suffix  (the edited final version)

When a pair is found and hasn't been analyzed yet:
  1. Extracts text from both versions and computes similarity.
  2. If similarity < THRESHOLD, calls Claude to identify patterns.
  3. Appends learned rules to write_like_user.md and/or resume_coverletter_guide.md.
  4. Marks the folder as analyzed (.feedback_applied) to avoid reprocessing.

How to use the feedback loop:
  1. The pipeline generates *_rough_draft.docx files.
  2. You review, copy the file (removing _rough_draft from the name), and edit.
  3. Next time the pipeline runs, this script picks up your edits and learns.

Usage:
    python learn_from_edits.py              # Process all unanalyzed pairs
    python learn_from_edits.py --dry-run    # Show what would be analyzed
    python learn_from_edits.py --reset Acme_Corp_2026-06-08  # Re-analyze a folder
"""

import argparse
import difflib
import json
import sys
from datetime import date
from pathlib import Path
from claude_utils import find_claude_exe, call_claude, extract_json

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

from docx import Document as DocxDocument

SCRIPT_DIR      = Path(__file__).resolve().parent
JOB_HUNT_DIR    = SCRIPT_DIR.parent
CLAUDE_CODE_DIR = JOB_HUNT_DIR.parent.parent
COMMON_DIR      = CLAUDE_CODE_DIR / "common"
CONFIG_DIR      = JOB_HUNT_DIR / "config"
APPS_DIR        = JOB_HUNT_DIR / "applications"

WRITE_LIKE_USER = COMMON_DIR / "write_like_user.md"
GUIDE_FILE      = CONFIG_DIR / "resume_coverletter_guide.md"
FEEDBACK_MARKER = ".feedback_applied"

SIMILARITY_THRESHOLD = 0.92  # below this ratio = significant edits worth learning from


# ── Document utilities ─────────────────────────────────────────────────────────

def extract_docx_text(path: Path) -> str:
    """Extract all paragraph text from a .docx file as a single string."""
    try:
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also pull text from table cells (resume uses a table for company/dates)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            paragraphs.append(p.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[Error reading {path.name}: {e}]"


def similarity(a: str, b: str) -> float:
    return difflib.SequenceMatcher(None, a, b).ratio()


# ── Pair discovery ─────────────────────────────────────────────────────────────

def find_edit_pairs() -> dict[Path, dict]:
    """
    Scan applications/ for unanalyzed (rough_draft, final) pairs.
    Returns {app_folder: {'resume': pair_dict, 'cover_letter': pair_dict}}.
    """
    result: dict[Path, dict] = {}

    if not APPS_DIR.exists():
        return result

    for app_folder in sorted(APPS_DIR.iterdir()):
        if not app_folder.is_dir():
            continue
        if (app_folder / FEEDBACK_MARKER).exists():
            continue

        folder_pairs: dict[str, dict] = {}
        for rough in app_folder.glob("*_rough_draft.docx"):
            final_name = rough.name.replace("_rough_draft.docx", ".docx")
            final = app_folder / final_name
            if final.exists():
                doc_type = "cover_letter" if "Cover_Letter" in rough.name else "resume"
                folder_pairs[doc_type] = {
                    "rough_draft": rough,
                    "final": final,
                    "app_folder": app_folder,
                }

        if folder_pairs:
            result[app_folder] = folder_pairs

    return result


# ── Prompt builder ─────────────────────────────────────────────────────────────

def build_analysis_prompt(
    app_name: str,
    pairs_with_text: dict,
    current_guide: str,
    current_voice: str,
) -> str:
    today = date.today().isoformat()
    doc_sections = []

    for doc_type, pair in pairs_with_text.items():
        label = "RESUME" if doc_type == "resume" else "COVER LETTER"
        doc_sections.append(f"""\
=== {label} — ROUGH DRAFT (LLM output) ===
{pair['rough_text']}

=== {label} — FINAL VERSION (edited) ===
{pair['final_text']}
""")

    return f"""\
You are analyzing the differences between LLM-generated job application documents
and the edited final versions. Your job is to extract reusable writing patterns
from the edits so future documents need fewer corrections.

Application folder: {app_name}
Analysis date: {today}

{"".join(doc_sections)}
=== CURRENT RESUME/COVER LETTER GUIDE (resume_coverletter_guide.md) ===
{current_guide}

=== CURRENT VOICE GUIDE (write_like_user.md) ===
{current_voice}

=== INSTRUCTIONS ===

Study what changed. Focus on extracting PATTERNS — rules that would improve
future documents across all applications, not observations specific to this one
company or role.

Ask yourself for each change:
  - Is this a style preference that should always apply? → add to write_like_user.md
  - Is this a structural or content rule? → add to resume_coverletter_guide.md
  - Is this just a company-specific word swap? → ignore it

Respond with ONLY a JSON object in a ```json code block:

```json
{{
  "significant_changes": true,
  "summary": "One sentence: the main theme of the edits",
  "voice_additions": "",
  "guide_additions": ""
}}
```

Rules for your additions:
- Specific > vague. "Don't write 'responsible for'" is weak. "Use 'built' instead of 'was responsible for building'" is strong.
- Don't add rules already covered in the current guides.
- Format additions under a dated heading, as bullet points:

  ---
  ## Learned — {today} ({app_name})
  - [rule]
  - [rule]

- Max 5 rules per document type. Quality over quantity.
- If changes don't meet the bar for a reusable rule (cosmetic, company-specific, or
  already covered), set "significant_changes": false and leave additions empty.
"""


# ── File updater ───────────────────────────────────────────────────────────────

def append_to_guide(path: Path, content: str) -> None:
    """Append a block of text to a guide file."""
    if not content.strip():
        return
    existing = path.read_text(encoding="utf-8")
    separator = "\n" if existing.endswith("\n") else "\n\n"
    path.write_text(existing + separator + content.strip() + "\n", encoding="utf-8")


# ── Main ───────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="Learn from document edits (feedback loop)")
    parser.add_argument("--dry-run", action="store_true",
                        help="Show what would be analyzed without making changes")
    parser.add_argument("--model", default="claude-opus-4-8")
    parser.add_argument("--verbose", action="store_true",
                        help="Print Claude's raw analysis output")
    parser.add_argument("--reset", metavar="FOLDER",
                        help="Remove the .feedback_applied marker from an application folder")
    args = parser.parse_args()

    # --reset: clear the marker so a folder can be re-analyzed
    if args.reset:
        marker = APPS_DIR / args.reset / FEEDBACK_MARKER
        if marker.exists():
            marker.unlink()
            print(f"Reset: {marker.parent.name}")
        else:
            print(f"No marker found in: {args.reset}")
        return

    pairs_by_folder = find_edit_pairs()

    if not pairs_by_folder:
        print("No unanalyzed edit pairs found.")
        print("To trigger: copy *_rough_draft.docx, remove '_rough_draft' from the name, edit and save.")
        return

    print(f"Found {len(pairs_by_folder)} folder(s) with edits to analyze.")

    if args.dry_run:
        for folder, by_type in pairs_by_folder.items():
            print(f"\n  {folder.name}:")
            for doc_type, pair in by_type.items():
                print(f"    {doc_type}: {pair['rough_draft'].name}")
                print(f"           → {pair['final'].name}")
        return

    try:
        claude_exe = find_claude_exe()
    except FileNotFoundError as e:
        print(f"[ERROR] {e}")
        sys.exit(1)

    current_guide = GUIDE_FILE.read_text(encoding="utf-8") if GUIDE_FILE.exists() else ""
    current_voice = WRITE_LIKE_USER.read_text(encoding="utf-8") if WRITE_LIKE_USER.exists() else ""

    for folder, by_type in pairs_by_folder.items():
        print(f"\n{'─' * 60}")
        print(f"Analyzing: {folder.name}")

        # Extract text, compute similarity, filter out insignificant changes
        significant_pairs: dict[str, dict] = {}
        for doc_type, pair in by_type.items():
            rough_text = extract_docx_text(pair["rough_draft"])
            final_text = extract_docx_text(pair["final"])
            ratio = similarity(rough_text, final_text)
            pair["rough_text"] = rough_text
            pair["final_text"] = final_text
            pair["similarity"] = ratio

            if ratio < SIMILARITY_THRESHOLD:
                print(f"  {doc_type}: {ratio:.0%} similar — changes detected, will analyze")
                significant_pairs[doc_type] = pair
            else:
                print(f"  {doc_type}: {ratio:.0%} similar — no significant edits, skipping")

        if not significant_pairs:
            print("  → All documents near-identical. Marking as analyzed.")
            (folder / FEEDBACK_MARKER).write_text(
                f"analyzed: {date.today().isoformat()}\nresult: no_significant_changes\n",
                encoding="utf-8",
            )
            continue

        print(f"  → Calling Claude to identify patterns...")
        prompt = build_analysis_prompt(folder.name, significant_pairs, current_guide, current_voice)

        try:
            raw_output = call_claude(prompt, claude_exe, args.model)
        except RuntimeError as e:
            print(f"  [ERROR] Claude call failed: {e}")
            continue

        if args.verbose:
            print("\n  --- Claude raw output ---")
            print(raw_output)
            print("  -------------------------\n")

        try:
            analysis = extract_json(raw_output)
        except (ValueError, json.JSONDecodeError) as e:
            print(f"  [ERROR] Could not parse Claude's response: {e}")
            print("  Tip: run with --verbose to see the raw output.")
            continue

        if not analysis.get("significant_changes", False):
            print("  → Claude: changes not pattern-worthy. No guide updates.")
            (folder / FEEDBACK_MARKER).write_text(
                f"analyzed: {date.today().isoformat()}\nresult: not_significant\n",
                encoding="utf-8",
            )
            continue

        print(f"  Summary: {analysis.get('summary', '')}")

        voice_additions = analysis.get("voice_additions", "").strip()
        guide_additions = analysis.get("guide_additions", "").strip()
        updated: list[str] = []

        if voice_additions and WRITE_LIKE_USER.exists():
            append_to_guide(WRITE_LIKE_USER, voice_additions)
            current_voice = WRITE_LIKE_USER.read_text(encoding="utf-8")
            updated.append("write_like_user.md")

        if guide_additions and GUIDE_FILE.exists():
            append_to_guide(GUIDE_FILE, guide_additions)
            current_guide = GUIDE_FILE.read_text(encoding="utf-8")
            updated.append("resume_coverletter_guide.md")

        if updated:
            print(f"  Updated: {', '.join(updated)}")
        else:
            print("  No additions needed (patterns already covered in guides).")

        (folder / FEEDBACK_MARKER).write_text(
            f"analyzed: {date.today().isoformat()}\n"
            f"summary: {analysis.get('summary', '')}\n"
            f"updated: {', '.join(updated) or 'none'}\n",
            encoding="utf-8",
        )

    print(f"\n{'─' * 60}")
    print("Done.")


if __name__ == "__main__":
    main()
