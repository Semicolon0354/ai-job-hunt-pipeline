#!/usr/bin/env python3
"""
Generate a tailored resume and cover letter from a job summary .txt file.
Uses a local Qwen model via Ollama to generate all content; python-docx handles formatting.

Usage:
    python generate_resume_coverletter.py --summary ../job_summaries/summary_Acme_2026-06-08.txt

Dependencies:
    pip install python-docx requests
"""
import argparse
import json
import re
import shutil
import sys
from datetime import date
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

# Windows console defaults to cp1252; ensure clean output in automated runs.
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

# ── Directory layout ──────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
JOB_HUNT_DIR = SCRIPT_DIR.parent
CLAUDE_CODE_DIR = JOB_HUNT_DIR.parent.parent
COMMON_DIR = CLAUDE_CODE_DIR / "common"
APPLICATIONS_DIR = JOB_HUNT_DIR / "applications"
CONFIG_DIR = JOB_HUNT_DIR / "config"

# ── Profile import ────────────────────────────────────────────────────────────
# All personal data lives in common/profile.py — edit there, not here.
sys.path.insert(0, str(COMMON_DIR))
import profile as user

CONTACT       = user.CONTACT
EDUCATION     = user.EDUCATION
SKILLS        = user.SKILLS
CERTS         = user.CERTS
INTERESTS     = user.INTERESTS
ORGANIZATIONS = user.ORGANIZATIONS

# ── LLM config ────────────────────────────────────────────────────────────────
DEFAULT_OLLAMA_URL = "http://localhost:11434"
DEFAULT_MODEL = "qwen3.5:9b"


# ── LLM interface ─────────────────────────────────────────────────────────────

def llm_chat(messages: list[dict], model: str, base_url: str) -> str:
    """Send a chat request to Ollama's native API; return the assistant reply as a string."""
    payload = {
        "model": model,
        "messages": messages,
        "stream": False,
        "think": False,
        "format": "json",
        "options": {
            "temperature": 0.3,
            "num_predict": 8192,
            "num_ctx": 32768,
        },
    }
    try:
        resp = requests.post(f"{base_url}/api/chat", json=payload, timeout=600)
        resp.raise_for_status()
    except requests.RequestException as e:
        sys.exit(f"ERROR: Ollama request failed — {e}\nIs Ollama running at {base_url}?")

    return resp.json()["message"]["content"]


def parse_json_response(raw: str, label: str) -> dict | list:
    """Parse JSON from a model response; extract from markdown fences if needed."""
    # Strip markdown code fences
    raw = re.sub(r"```(?:json)?\s*", "", raw).strip().rstrip("`").strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            try:
                return json.loads(m.group())
            except json.JSONDecodeError:
                pass
        raise ValueError(f"Could not parse {label} JSON from model response.\nRaw output:\n{raw[:600]}")


# ── Summary parsing ───────────────────────────────────────────────────────────

def parse_summary(text: str) -> dict:
    """Extract company, role, score, and date from a job summary .txt."""
    data = {"company": "", "role": "", "score": "", "date": ""}
    for line in text.splitlines():
        low = line.lower().strip()
        # Header line: JOB SUMMARY - Company Name | Job Title
        m = re.search(r"JOB SUMMARY - (.+?) \| (.+)", line, re.IGNORECASE)
        if m:
            data["company"] = m.group(1).strip()
            data["role"] = m.group(2).strip()
        elif low.startswith("score:"):
            data["score"] = line.split(":", 1)[1].strip()
        elif low.startswith("date:"):
            data["date"] = line.split(":", 1)[1].strip()
    return data


# ── LLM prompts ───────────────────────────────────────────────────────────────

def generate_resume_content(
    summary_text: str, profile_text: str, guide_text: str, model: str, base_url: str
) -> dict:
    """Ask the LLM to produce tailored resume sections as structured JSON."""
    system = (
        "You are a professional resume writer. "
        f"Tailor {CONTACT['name']}'s resume to the job posting provided. "
        "Output ONLY valid JSON — no prose, no markdown fences, no commentary.\n\n"
        "TAILORING GUIDE (follow these rules exactly):\n"
        + guide_text
    )
    user = f"""Generate tailored resume content for the job summary below.
Follow all rules in the tailoring guide above.

Return a JSON object with EXACTLY this structure:
{{
  "professional_summary": "2-3 sentence summary targeted to this specific role",
  "core_competencies": ["skill1", "skill2"],
  "main_experience": [
    {{
      "company": "exact company name from profile",
      "title": "exact job title from profile",
      "dates": "exact dates from profile",
      "bullets": ["bullet1", "bullet2", "bullet3", "bullet4"]
    }}
  ],
  "additional_experience": [
    {{
      "company": "exact company name from profile",
      "title": "exact job title from profile",
      "dates": "exact dates from profile",
      "bullets": ["bullet1", "bullet2"]
    }}
  ],
  "education": [
    {{
      "degree": "exact degree from profile",
      "school": "exact school from profile",
      "date": "exact date from profile",
      "coursework": ["3-5 most relevant courses for this role"]
    }}
  ],
  "skills": [
    {{
      "category": "exact category name from profile",
      "items": ["skill1", "skill2"]
    }}
  ]
}}

STRICT RULES:
- Every role in main_experience must have EXACTLY 4 bullets — all roles, not just the first
- Every role in additional_experience must have EXACTLY 2 bullets
- NEVER move a role from additional_experience to main_experience to meet the bullet count
- Role placement is fixed by profile rules — bullet count does not change it
- Bullets must be selected from the profile's existing bullets for each role. Minor wording adjustments to incorporate ATS keywords are allowed, but the core accomplishment and message must remain intact. Do not invent bullets.
- education: keep degree, school, and date exactly as in the profile; select 3-5 coursework items from the profile's coursework list — use course titles verbatim, do not invent or rename any course
- skills: include ALL categories and items from the profile; reorder items within each category to put job-relevant tools first; never add or remove skills, never add Tableau

CANDIDATE PROFILE:
{profile_text}

JOB SUMMARY:
{summary_text}"""

    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]
    for attempt in range(1, 4):
        raw = llm_chat(messages, model, base_url)
        try:
            return parse_json_response(raw, "resume content")
        except ValueError as e:
            if attempt < 3:
                print(f"  [retry {attempt}/3] resume JSON parse failed, retrying...")
            else:
                sys.exit(f"ERROR: {e}")


def generate_cover_letter_paragraphs(
    summary_text: str,
    profile_text: str,
    voice_text: str,
    guide_text: str,
    company: str,
    role: str,
    model: str,
    base_url: str,
) -> list[str]:
    """Ask the LLM to write a 4-paragraph cover letter in the user's voice."""
    system = (
        f"You are writing in {CONTACT['name']}'s voice. Apply EVERY pattern in the writing guide below. "
        "Do not revert to generic AI writing style. "
        "Output ONLY valid JSON — no prose, no markdown fences.\n\n"
        "=== VOICE AND WRITING STYLE (write_like_user.md) ===\n"
        + voice_text
        + "\n\n=== COVER LETTER RULES (resume_coverletter_guide.md) ===\n"
        + guide_text
    )
    user = f"""Write a 4-paragraph cover letter for {CONTACT['name']} applying for the {role} position at {company}.
Follow the cover letter structure and voice rules in the guides above.

Return ONLY this JSON structure:
{{
  "paragraphs": [
    "Paragraph 1 text",
    "Paragraph 2 text",
    "Paragraph 3 text",
    "Paragraph 4 text"
  ]
}}

CANDIDATE PROFILE:
{profile_text}

JOB SUMMARY:
{summary_text}"""

    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]
    for attempt in range(1, 4):
        raw = llm_chat(messages, model, base_url)
        try:
            result = parse_json_response(raw, "cover letter")
            if isinstance(result, dict):
                return result.get("paragraphs", [])
            return result
        except ValueError as e:
            if attempt < 3:
                print(f"  [retry {attempt}/3] cover letter JSON parse failed, retrying...")
            else:
                sys.exit(f"ERROR: {e}")


# ── Document builders ─────────────────────────────────────────────────────────

def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    # Underline separator via bottom border would require XML; keep simple for ATS safety


def _add_body_paragraph(doc: Document, text: str, size: float = 10.5) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)


def _add_experience_entry(doc: Document, company: str, title: str, dates: str, bullets: list[str]) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.allow_autofit = False
    tbl.columns[0].width = Inches(5.25)
    tbl.columns[1].width = Inches(2.25)

    left = tbl.cell(0, 0).paragraphs[0]
    r = left.add_run(f"{title}  |  {company}")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)

    right = tbl.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = right.add_run(dates)
    r2.italic = True
    r2.font.name = "Calibri"
    r2.font.size = Pt(10.5)

    for bullet in bullets:
        _add_bullet(doc, bullet)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_resume(content: dict) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(CONTACT["pro_name"])
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(2)

    # Contact line
    contact_line = (
        f"{CONTACT['email']} | {CONTACT['phone']} | {CONTACT['location']}\n"
        f"{CONTACT['linkedin']}  |  {CONTACT['github']}"
    )
        
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(contact_line)
    r2.font.name = "Calibri"
    r2.font.size = Pt(9.5)
    p2.paragraph_format.space_after = Pt(8)

    # Professional Summary
    _add_section_heading(doc, "Professional Summary")
    _add_body_paragraph(doc, content.get("professional_summary", ""))

    # Core Competencies
    _add_section_heading(doc, "Core Competencies")
    competencies = content.get("core_competencies", [])
    _add_body_paragraph(doc, "  •  ".join(competencies))

    # Professional Experience
    _add_section_heading(doc, "Professional Experience")
    for exp in content.get("main_experience", []):
        _add_experience_entry(doc, exp.get("company", ""), exp.get("title", ""), exp.get("dates", ""), exp.get("bullets", []))

    # Additional Experience (Paramedic for non-healthcare roles)
    additional = content.get("additional_experience", [])
    if additional:
        _add_section_heading(doc, "Additional Experience")
        for exp in additional:
            _add_experience_entry(doc, exp.get("company", ""), exp.get("title", ""), exp.get("dates", ""), exp.get("bullets", []))

    # Education — use LLM-tailored coursework, fall back to full profile list
    _add_section_heading(doc, "Education")
    edu_list = content.get("education") or EDUCATION
    for edu in edu_list:
        p = doc.add_paragraph()
        r = p.add_run(f"{edu['degree']}  |  {edu['school']}  |  {edu['date']}")
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(1)
        cw = doc.add_paragraph()
        cw.add_run(f"Relevant Coursework: {', '.join(edu['coursework'])}")
        for r in cw.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(10)
        cw.paragraph_format.space_after = Pt(4)

    # Technical Skills — use LLM-reordered skills, fall back to full profile list
    _add_section_heading(doc, "Technical Skills")
    skills_list = content.get("skills")
    if skills_list and isinstance(skills_list, list):
        for entry in skills_list:
            category = entry.get("category", "")
            items = entry.get("items", [])
            p = doc.add_paragraph()
            r_label = p.add_run(f"{category}: ")
            r_label.bold = True
            r_label.font.name = "Calibri"
            r_label.font.size = Pt(10.5)
            r_items = p.add_run(", ".join(items))
            r_items.font.name = "Calibri"
            r_items.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(2)
    else:
        for category, items in SKILLS.items():
            p = doc.add_paragraph()
            r_label = p.add_run(f"{category}: ")
            r_label.bold = True
            r_label.font.name = "Calibri"
            r_label.font.size = Pt(10.5)
            r_items = p.add_run(", ".join(items))
            r_items.font.name = "Calibri"
            r_items.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(2)

    # Interests & Organizations
    _add_section_heading(doc, "Interests & Organizations")
    interests_str = "  |  ".join(
        f"{activity}: {', '.join(disciplines)}" for activity, disciplines in INTERESTS.items()
    )
    _add_body_paragraph(doc, interests_str)
    _add_body_paragraph(doc, "  |  ".join(ORGANIZATIONS))

    return doc


def build_cover_letter(paragraphs: list[str], company: str, role: str, date_str: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    # Header
    p = doc.add_paragraph()
    r = p.add_run(CONTACT["pro_name"])
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(1)

    header_line = (
        f"{CONTACT['email']} | {CONTACT['phone']} | {CONTACT['location']}\n"
        f"{CONTACT['linkedin']}  |  {CONTACT['github']}"
    )
    _add_body_paragraph(doc, header_line)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _add_body_paragraph(doc, date_str)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _add_body_paragraph(doc, f"RE: {role}")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _add_body_paragraph(doc, "Dear Hiring Manager,")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    for para in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(para)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(6)

    _add_body_paragraph(doc, "Respectfully,")
    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    p = doc.add_paragraph()
    r = p.add_run(CONTACT["name"])
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)

    return doc


# ── Filename helpers ──────────────────────────────────────────────────────────

def sanitize(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9_\-]", "_", s).strip("_")


# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="Generate tailored resume and cover letter from a job summary")
    parser.add_argument("--summary", required=True, help="Path to job summary .txt file")
    parser.add_argument("--model", default=DEFAULT_MODEL, help=f"Ollama model name (default: {DEFAULT_MODEL})")
    parser.add_argument("--ollama-url", default=DEFAULT_OLLAMA_URL, help=f"Ollama base URL (default: {DEFAULT_OLLAMA_URL})")
    parser.add_argument("--date", default=date.today().strftime("%B %d, %Y"), help="Cover letter date")
    args = parser.parse_args()

    summary_path = Path(args.summary).resolve()
    if not summary_path.is_file():
        sys.exit(f"ERROR: Summary file not found: {summary_path}")

    # Read all inputs
    summary_text = summary_path.read_text(encoding="utf-8", errors="replace")
    profile_text = user.to_markdown()
    voice_text = (COMMON_DIR / "write_like_user.md").read_text(encoding="utf-8")
    guide_text = (CONFIG_DIR / "resume_coverletter_guide.md").read_text(encoding="utf-8")

    meta = parse_summary(summary_text)
    company = meta["company"] or "Unknown_Company"
    role = meta["role"] or "Unknown_Role"
    doc_date = meta["date"] or args.date

    print(f"Generating documents for: {role} @ {company}")
    print(f"Model: {args.model}")

    # LLM calls
    print("  [1/2] Generating resume content...")
    resume_content = generate_resume_content(summary_text, profile_text, guide_text, args.model, args.ollama_url)

    print("  [2/2] Generating cover letter...")
    cl_paragraphs = generate_cover_letter_paragraphs(
        summary_text, profile_text, voice_text, guide_text, company, role, args.model, args.ollama_url
    )

    # Output directory — use the date from the summary file, not today's date
    folder_date = meta["date"] or date.today().strftime("%Y-%m-%d")
    folder_name = f"{sanitize(company)}_{folder_date}"
    out_dir = APPLICATIONS_DIR / folder_name
    out_dir.mkdir(parents=True, exist_ok=True)

    last_name = sanitize(CONTACT["name"].split()[-1])

    # Save resume (marked as rough draft — copy and edit to trigger the feedback loop)
    resume_doc = build_resume(resume_content)
    resume_path = out_dir / f"{last_name}_Resume_{sanitize(company)}_{sanitize(role)}_rough_draft.docx"
    resume_doc.save(resume_path)
    print(f"\nSaved resume:       {resume_path}")

    # Save cover letter (same rough draft convention)
    cl_doc = build_cover_letter(cl_paragraphs, company, role, doc_date)
    cl_path = out_dir / f"{last_name}_Cover_Letter_{sanitize(company)}_{sanitize(role)}_rough_draft.docx"
    cl_doc.save(cl_path)
    print(f"Saved cover letter: {cl_path}")

    # Move summary file into application folder
    dest_summary = out_dir / summary_path.name
    shutil.move(str(summary_path), dest_summary)
    print(f"Moved summary:      {dest_summary}")

    print(f"\nDone — all files in: {out_dir}")


if __name__ == "__main__":
    main()
